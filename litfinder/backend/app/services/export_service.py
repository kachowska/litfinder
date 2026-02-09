"""
Bibliography Export Service
Export bibliographies to Word (.docx), BibTeX, and RIS formats.
"""
from typing import List, Optional
from io import BytesIO
from datetime import datetime
import re

from app.services.gost_formatter import (
    BibliographyEntry, 
    gost_formatter, 
    article_to_bibliography_entry,
    SourceType
)


class ExportService:
    """Export bibliographies to various formats."""
    
    def export_to_text(
        self, 
        entries: List[BibliographyEntry],
        sort_by: str = "author"
    ) -> str:
        """Export to plain text with GOST formatting."""
        formatted = gost_formatter.format_list(entries, sort_by)
        return "\n".join(formatted)
    
    def export_to_bibtex(self, entries: List[BibliographyEntry]) -> str:
        """
        Export to BibTeX format.
        
        Example:
        @article{author2024,
            author = {Author, I. O.},
            title = {Title},
            journal = {Journal},
            year = {2024},
            volume = {1},
            pages = {1--10}
        }
        """
        bibtex_entries = []
        
        for entry in entries:
            bibtex_entries.append(self._entry_to_bibtex(entry))
        
        return "\n\n".join(bibtex_entries)
    
    def export_to_ris(self, entries: List[BibliographyEntry]) -> str:
        """
        Export to RIS format (for EndNote, Zotero, Mendeley).
        
        Example:
        TY  - JOUR
        AU  - Author, I. O.
        TI  - Title
        JO  - Journal
        PY  - 2024
        VL  - 1
        SP  - 1
        EP  - 10
        ER  -
        """
        ris_entries = []
        
        for entry in entries:
            ris_entries.append(self._entry_to_ris(entry))
        
        return "\n".join(ris_entries)
    
    def export_to_docx(
        self, 
        entries: List[BibliographyEntry],
        title: str = "Список литературы",
        sort_by: str = "author"
    ) -> bytes:
        """
        Export to Word document (.docx).
        
        Returns:
            Bytes of the .docx file
        """
        try:
            from docx import Document
            from docx.shared import Pt, Cm
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            # Return simple text if python-docx not available
            return self.export_to_text(entries, sort_by).encode('utf-8')
        
        doc = Document()
        
        # Set document margins
        section = doc.sections[0]
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        
        # Add title
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Format entries
        formatted = gost_formatter.format_list(entries, sort_by)
        
        # Add each entry
        for text in formatted:
            para = doc.add_paragraph(text)
            para.paragraph_format.first_line_indent = Cm(-1)
            para.paragraph_format.left_indent = Cm(1)
            para.paragraph_format.space_after = Pt(6)
            
            # Set font
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.read()
    
    def _entry_to_bibtex(self, entry: BibliographyEntry) -> str:
        """Convert entry to BibTeX format."""
        # Determine entry type
        type_map = {
            SourceType.ARTICLE: "article",
            SourceType.BOOK: "book",
            SourceType.CONFERENCE: "inproceedings",
            SourceType.THESIS: "phdthesis",
            SourceType.ELECTRONIC: "misc"
        }
        entry_type = type_map.get(entry.source_type, "article")
        
        # Create citation key
        first_author = entry.authors[0].last_name if entry.authors else "unknown"
        year = entry.year or "nd"
        cite_key = self._make_cite_key(first_author, year, entry.title)
        
        # Build fields
        fields = []
        
        # Authors
        if entry.authors:
            author_str = " and ".join(
                f"{a.last_name}, {a.initials}" for a in entry.authors
            )
            fields.append(f'    author = {{{author_str}}}')
        
        # Title
        fields.append(f'    title = {{{entry.title}}}')
        
        # Type-specific fields
        if entry.source_type == SourceType.ARTICLE:
            if entry.journal_name:
                fields.append(f'    journal = {{{entry.journal_name}}}')
        elif entry.source_type == SourceType.BOOK:
            if entry.publisher:
                fields.append(f'    publisher = {{{entry.publisher}}}')
            if entry.city:
                fields.append(f'    address = {{{entry.city}}}')
        elif entry.source_type == SourceType.CONFERENCE:
            if entry.conference_name:
                fields.append(f'    booktitle = {{{entry.conference_name}}}')
        
        # Common fields
        if entry.year:
            fields.append(f'    year = {{{entry.year}}}')
        if entry.volume:
            fields.append(f'    volume = {{{entry.volume}}}')
        if entry.issue:
            fields.append(f'    number = {{{entry.issue}}}')
        if entry.pages:
            # Convert "1-10" to "1--10"
            pages = entry.pages.replace("-", "--")
            fields.append(f'    pages = {{{pages}}}')
        if entry.doi:
            fields.append(f'    doi = {{{entry.doi}}}')
        if entry.url:
            fields.append(f'    url = {{{entry.url}}}')
        
        return f"@{entry_type}{{{cite_key},\n" + ",\n".join(fields) + "\n}"
    
    def _entry_to_ris(self, entry: BibliographyEntry) -> str:
        """Convert entry to RIS format."""
        # Determine entry type
        type_map = {
            SourceType.ARTICLE: "JOUR",
            SourceType.BOOK: "BOOK",
            SourceType.CONFERENCE: "CONF",
            SourceType.THESIS: "THES",
            SourceType.ELECTRONIC: "ELEC"
        }
        ris_type = type_map.get(entry.source_type, "JOUR")
        
        lines = [f"TY  - {ris_type}"]
        
        # Authors
        for author in entry.authors:
            lines.append(f"AU  - {author.last_name}, {author.initials}")
        
        # Title
        lines.append(f"TI  - {entry.title}")
        
        # Type-specific fields
        if entry.journal_name:
            lines.append(f"JO  - {entry.journal_name}")
        if entry.publisher:
            lines.append(f"PB  - {entry.publisher}")
        if entry.city:
            lines.append(f"CY  - {entry.city}")
        
        # Common fields
        if entry.year:
            lines.append(f"PY  - {entry.year}")
        if entry.volume:
            lines.append(f"VL  - {entry.volume}")
        if entry.issue:
            lines.append(f"IS  - {entry.issue}")
        
        # Pages
        if entry.pages:
            if "-" in entry.pages:
                start, end = entry.pages.split("-", 1)
                lines.append(f"SP  - {start.strip()}")
                lines.append(f"EP  - {end.strip()}")
            else:
                lines.append(f"SP  - {entry.pages}")
        
        if entry.doi:
            lines.append(f"DO  - {entry.doi}")
        if entry.url:
            lines.append(f"UR  - {entry.url}")
        
        lines.append("ER  -")
        
        return "\n".join(lines)
    
    def _make_cite_key(self, author: str, year, title: str) -> str:
        """Generate a unique citation key."""
        # Clean author name
        author_clean = re.sub(r'[^a-zA-Zа-яА-Я]', '', author)[:10]
        
        # Get first word of title
        title_word = re.sub(r'[^a-zA-Zа-яА-Я]', '', title.split()[0]) if title else "untitled"
        title_word = title_word[:10]
        
        return f"{author_clean}{year}{title_word}".lower()


# --- API helpers ---

def export_articles(
    articles: List[dict],
    format: str = "gost",
    sort_by: str = "author"
) -> dict:
    """
    Export list of articles to specified format.
    
    Args:
        articles: List of article dicts
        format: Export format (gost, bibtex, ris, docx)
        sort_by: Sort order (author, year, title)
        
    Returns:
        Dict with format, content, and filename
    """
    service = ExportService()
    
    # Convert articles to entries
    entries = [article_to_bibliography_entry(a) for a in articles]
    
    if format == "gost" or format == "text":
        content = service.export_to_text(entries, sort_by)
        return {
            "format": "text",
            "content": content,
            "filename": "bibliography.txt",
            "mime_type": "text/plain"
        }
    
    elif format == "bibtex":
        content = service.export_to_bibtex(entries)
        return {
            "format": "bibtex",
            "content": content,
            "filename": "bibliography.bib",
            "mime_type": "application/x-bibtex"
        }
    
    elif format == "ris":
        content = service.export_to_ris(entries)
        return {
            "format": "ris",
            "content": content,
            "filename": "bibliography.ris",
            "mime_type": "application/x-research-info-systems"
        }
    
    elif format == "docx" or format == "word":
        content_bytes = service.export_to_docx(entries, sort_by=sort_by)
        import base64
        return {
            "format": "docx",
            "content": base64.b64encode(content_bytes).decode('utf-8'),
            "filename": "bibliography.docx",
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "is_binary": True
        }
    
    else:
        return {
            "format": "error",
            "content": f"Unsupported format: {format}",
            "filename": None,
            "mime_type": None
        }


# --- Singleton instance ---
export_service = ExportService()
